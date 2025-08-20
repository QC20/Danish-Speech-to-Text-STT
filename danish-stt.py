"""
Header Information
Date: 20 August 2025
Author: Jonas Kjeldmand Jensen
Affiliation: Technical University of Denmark, Section for Strategy and Leadership Development
Description: This script performs enhanced transcription of research interviews using OpenAI Whisper Large-v3.
dependencies: pip install openai-whisper python-docx
git clone: https://github.com/QC20/Danish-Speech-to-Text-STT.git
"""

import whisper
import os
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import datetime

def get_desktop_path():
    """Get desktop path that works on any operating system"""
    if os.name == 'nt':  # Windows
        return os.path.join(os.path.expanduser('~'), 'Desktop')
    else:  # macOS and Linux
        return os.path.join(os.path.expanduser('~'), 'Desktop')

def create_project_folder(base_name="interview_transcript"):
    """Create a unique project folder on desktop"""
    desktop = get_desktop_path()
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    folder_name = f"{base_name}_{timestamp}"
    folder_path = os.path.join(desktop, folder_name)
    
    os.makedirs(folder_path, exist_ok=True)
    print(f"Created project folder: {folder_path}")
    return folder_path

def format_timestamp(seconds):
    """Convert seconds to MM:SS format"""
    minutes = int(seconds // 60)
    seconds = int(seconds % 60)
    return f"{minutes:02d}:{seconds:02d}"

def format_duration(seconds):
    """Convert seconds to HH:MM:SS format for total duration"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = int(seconds % 60)
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    else:
        return f"{minutes:02d}:{seconds:02d}"

def detect_speaker_changes(segments, silence_threshold=1.5, min_speaker_time=3.0):
    """Speaker change detection with minimum speaker time constraint"""
    speakers = []
    current_speaker = "INTERVIEWER"
    last_speaker_change_time = 0
    
    for i, segment in enumerate(segments):
        if i == 0:
            speakers.append(current_speaker)
        else:
            prev_end = segments[i-1]['end']
            current_start = segment['start']
            gap = current_start - prev_end
            time_since_last_change = current_start - last_speaker_change_time
            
            # Switch speakers if there's a significant pause AND enough time has passed
            if gap > silence_threshold and time_since_last_change > min_speaker_time:
                current_speaker = "PARTICIPANT" if current_speaker == "INTERVIEWER" else "INTERVIEWER"
                last_speaker_change_time = current_start
            
            speakers.append(current_speaker)
    
    return speakers

def add_content_markers(text):
    """Add content markers for unclear speech, pauses, etc."""
    # Simple heuristics for content markers
    if len(text.strip()) < 3:
        return "[UNCLEAR]"
    
    # Add more sophisticated detection here
    markers = text
    
    # Replace common whisper artifacts
    markers = markers.replace("...", "[PAUSE]")
    markers = markers.replace("  ", " ")
    
    return markers.strip()

def create_academic_word_document(transcript_data, metadata, output_path):
    """Create an academic-style Word document"""
    print("Creating academic-style Word document...")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('INTERVIEW TRANSCRIPT', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Metadata section
    doc.add_heading('Study Information', level=1)
    
    # Create metadata table
    table = doc.add_table(rows=12, cols=2)
    table.style = 'Table Grid'
    
    metadata_items = [
        ('Research Study ID:', '[Study ID Here]'),
        ('Interviewer:', metadata['interviewer']),
        ('Participant:', metadata['participant']),
        ('Date:', metadata['date']),
        ('Duration:', metadata['duration']),
        ('Location:', '[Location Here]'),
        ('Audio Quality:', metadata['audio_quality']),
        ('Transcription Method:', 'OpenAI Whisper Large-v3 (Automated) + Manual Review'),
        ('Total Segments:', str(metadata['total_segments'])),
        ('Average Confidence:', f"{metadata['avg_confidence']} (Higher is better)"),
        ('Context Provided:', f"Yes - {metadata['interview_type']} context"),
        ('Quality Parameters:', 'Enhanced (beam search, best-of-5, context-aware)')
    ]
    
    for i, (key, value) in enumerate(metadata_items):
        table.cell(i, 0).text = key
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
        table.cell(i, 1).text = value
    
    # Transcription notes
    doc.add_heading('Transcription Notes', level=1)
    notes_para = doc.add_paragraph()
    notes_para.add_run("• Automated transcription using OpenAI Whisper Large-v3 model\n")
    notes_para.add_run(f"• Language: {metadata['language']} (explicitly specified)\n") 
    notes_para.add_run("• Enhanced quality parameters: beam search, best-of-5, context-aware\n")
    notes_para.add_run(f"• Context provided: {metadata['interview_type']} setting\n")
    notes_para.add_run("• Speaker identification based on intelligent pause detection\n")
    notes_para.add_run("• Timestamps indicate start time of each utterance\n")
    notes_para.add_run("• [PAUSE], [UNCLEAR] markers added where detected\n")
    notes_para.add_run("• Low-confidence segments filtered for quality\n")
    notes_para.add_run("• Manual review recommended for final accuracy verification\n")
    
    # Participants information
    doc.add_heading('Participants', level=1)
    participants_para = doc.add_paragraph()
    participants_para.add_run("INTERVIEWER: ").bold = True
    participants_para.add_run(f"{metadata['interviewer']}\n")
    participants_para.add_run("PARTICIPANT: ").bold = True  
    participants_para.add_run(f"{metadata['participant']}\n")
    
    # Enhanced statistics
    doc.add_heading('Interview Statistics', level=1)
    
    # Create statistics table for better formatting
    stats_table = doc.add_table(rows=15, cols=2)
    stats_table.style = 'Light Grid Accent 1'
    
    # Calculate speaker-specific statistics
    interviewer_turns = len([t for t in transcript_data if t['speaker'] == 'INTERVIEWER'])
    participant_turns = len([t for t in transcript_data if t['speaker'] == 'PARTICIPANT'])
    
    stats_items = [
        ('Total Turns:', str(len(transcript_data))),
        ('Interviewer Turns:', f"{interviewer_turns} ({(interviewer_turns/len(transcript_data)*100):.1f}%)"),
        ('Participant Turns:', f"{participant_turns} ({(participant_turns/len(transcript_data)*100):.1f}%)"),
        ('Total Words:', str(metadata['total_words'])),
        ('Avg Words/Segment:', str(metadata['avg_words_per_segment'])),
        ('Speech Rate:', f"{metadata['total_words']/(float(metadata['duration'].split(':')[0])*60 + float(metadata['duration'].split(':')[1])):.1f} words/min"),
        ('Avg Segment Length:', f"{metadata['avg_segment_length']}s"),
        ('Longest Segment:', f"{metadata['longest_segment']}s"),
        ('Shortest Segment:', f"{metadata['shortest_segment']}s"),
        ('Total Pauses:', str(metadata['total_pauses'])),
        ('Avg Pause Duration:', f"{metadata['avg_pause_duration']}s"),
        ('Longest Pause:', f"{metadata['longest_pause']}s"),
        ('Speech vs Silence:', f"{metadata['speech_vs_pause_ratio']}% speech"),
        ('High Confidence Segs:', f"{metadata['high_confidence_pct']}%"),
        ('Overall Confidence:', f"{metadata['avg_confidence']} (scale: -3 to 0)")
    ]
    
    for i, (key, value) in enumerate(stats_items):
        stats_table.cell(i, 0).text = key
        stats_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        stats_table.cell(i, 1).text = value
    
    # Main transcript
    doc.add_page_break()
    doc.add_heading('TRANSCRIPT', level=1)
    
    # Add turn numbers
    current_speaker = None
    turn_number = 0
    
    for segment in transcript_data:
        if current_speaker != segment['speaker']:
            if current_speaker is not None:
                doc.add_paragraph("")  # Add space between speakers
            current_speaker = segment['speaker']
            turn_number += 1
        
        # Create turn paragraph
        p = doc.add_paragraph()
        
        # Add turn number and timestamp
        turn_info = f"Turn {turn_number:2d} [{segment['start_time']} - {segment['end_time']}] "
        run1 = p.add_run(turn_info)
        run1.bold = True
        run1.font.size = run1.font.size
        
        # Add speaker
        run2 = p.add_run(f"{segment['speaker']}: ")
        run2.bold = True
        if segment['speaker'] == 'INTERVIEWER':
            run2.font.color.rgb = None  # Default color
        else:
            run2.font.color.rgb = None  # You can set different colors here
        
        # Add text with content markers
        marked_text = add_content_markers(segment['text'])
        p.add_run(marked_text)
    
    # Footer with page numbers
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Interview Transcript - Page "
    
    doc.save(output_path)
    print(f"Academic Word document saved as: {output_path}")

def get_audio_quality_estimate(result):
    """Estimate audio quality based on Whisper's internal metrics"""
    # Simple heuristic based on segment characteristics
    segments = result.get('segments', [])
    if not segments:
        return "Unknown"
    
    avg_confidence = sum(seg.get('avg_logprob', -1) for seg in segments) / len(segments)
    
    if avg_confidence > -0.5:
        return "High"
    elif avg_confidence > -1.0:
        return "Medium"
    else:
        return "Low"

def main():
    # Configuration - UPDATE THESE FOR EACH INTERVIEW
    # =================================================
    # AUDIO FILE SETTINGS
    audio_file_path = "/path/to/your/audio/file.mp3"  # UPDATE: Path to your audio file
    
    # PARTICIPANT INFORMATION
    interviewer_name = "Interviewer ([Your Name/Role])"  # UPDATE: Interviewer identification
    participant_name = "Participant ([Participant Code/Role])"  # UPDATE: Participant identification
    
    # LANGUAGE SETTINGS
    language_code = "en"  # UPDATE: Language code (e.g., "en", "da", "de", "fr", etc.)
    
    # INTERVIEW CONTEXT DESCRIPTION
    # UPDATE: Provide detailed context about your interview to improve transcription accuracy
    interview_context = (
        "This is a [DESCRIPTION OF INTERVIEW TYPE] conducted in [LANGUAGE]. "
        "The interview focuses on [MAIN TOPIC/RESEARCH AREA]. "
        "It is conducted between [INTERVIEWER ROLE/DESCRIPTION] and [PARTICIPANT ROLE/DESCRIPTION]. "
        "The participant [BRIEF PARTICIPANT BACKGROUND/CONTEXT]. "
        "[ADDITIONAL RELEVANT CONTEXT THAT WOULD HELP THE TRANSCRIPTION MODEL UNDERSTAND THE DOMAIN/SETTING]. "
    )
    
    # EXAMPLE OF FILLED CONTEXT (Danish research interview):
    # interview_context = (
    #     "This is a Danish research interview about AI and leadership. "
    #     "The interview focuses on how AI technology can potentially change leadership roles "
    #     "and its broader influence on organizations. "
    #     "It is conducted between a researcher and a company leader. "
    #     "The participant is a leader in a software company that develops survey tools "
    #     "for organizations, typically used for employee satisfaction surveys. "
    #     "The company is small with under 20 employees. "
    # )
    # =================================================

    # Enhanced Whisper parameters with all quality improvements
    whisper_params = {
        "language": language_code,  # Use configured language
        "task": "transcribe",
        "word_timestamps": True,
        "verbose": True,
        #"hallucination_silence_threshold": 1.0, # Only work with large-v3
        
        # Context and accuracy improvements
        "condition_on_previous_text": True,  # Use previous text for context
        "initial_prompt": interview_context,  # Provide context to model
        
        # Quality filtering
        "compression_ratio_threshold": 2.4,  # Skip overly compressed segments
        "logprob_threshold": -1.0,  # Skip low-confidence segments  
        "no_speech_threshold": 0.6,  # Better silence detection
        
        # Decoding improvements
        "temperature": 0.0,  # More deterministic output (less random)
        "best_of": 5,  # Try 5 different decoding attempts, use best
        "beam_size": 5,  # Beam search for better accuracy
        
        # Token control
        "suppress_tokens": [-1],  # Suppress unwanted tokens
        "prepend_punctuations": "\"'¿([{-",  # Better punctuation handling
        "append_punctuations": "\"'.。,，!！?？:：)]}、",  # Better punctuation handling
    }
    
    # Check if file exists
    if not os.path.exists(audio_file_path):
        print(f"Error: Audio file not found at {audio_file_path}")
        print("Please update the audio_file_path variable with the correct path to your audio file.")
        return
    
    # Create project folder
    project_folder = create_project_folder("academic_interview")
    
    # Copy audio file to project folder
    audio_filename = os.path.basename(audio_file_path)
    audio_copy_path = os.path.join(project_folder, audio_filename)
    shutil.copy2(audio_file_path, audio_copy_path)
    print(f"Audio file copied to project folder")
    
    print("Loading Whisper model (large-v3 for maximum accuracy)...")
    # Use largest model for best accuracy
    whisper_model = whisper.load_model("large-v3")  # Maximum accuracy
    
    print("Transcribing audio with enhanced settings...")
    print("This may take longer but will provide better accuracy...\n")
    
    # Transcribe with enhanced parameters
    result = whisper_model.transcribe(audio_file_path, **whisper_params)
    
    print("\nTranscription complete! Processing segments...")
    
    # Get enhanced metadata with statistics
    audio_duration = result.get('duration', 0)
    segments = result["segments"]
    
    # Calculate speaking statistics
    total_words = sum(len(seg['text'].split()) for seg in segments)
    avg_words_per_segment = total_words / len(segments) if segments else 0
    
    # Calculate confidence statistics
    confidences = [seg.get('avg_logprob', -2) for seg in segments]
    avg_confidence = sum(confidences) / len(confidences) if confidences else -2
    high_confidence_segments = len([c for c in confidences if c > -0.5])
    medium_confidence_segments = len([c for c in confidences if -1.0 <= c <= -0.5])
    low_confidence_segments = len([c for c in confidences if c < -1.0])
    
    # Calculate timing statistics
    segment_lengths = [seg['end'] - seg['start'] for seg in segments]
    avg_segment_length = sum(segment_lengths) / len(segment_lengths) if segment_lengths else 0
    longest_segment = max(segment_lengths) if segment_lengths else 0
    shortest_segment = min(segment_lengths) if segment_lengths else 0
    
    # Calculate pause statistics
    pauses = []
    for i in range(1, len(segments)):
        pause_duration = segments[i]['start'] - segments[i-1]['end']
        if pause_duration > 0:
            pauses.append(pause_duration)
    
    avg_pause_duration = sum(pauses) / len(pauses) if pauses else 0
    longest_pause = max(pauses) if pauses else 0
    total_pause_time = sum(pauses)
    
    # Get language name for display
    language_names = {
        "en": "English", "da": "Danish", "de": "German", "fr": "French", 
        "es": "Spanish", "it": "Italian", "pt": "Portuguese", "nl": "Dutch",
        "sv": "Swedish", "no": "Norwegian", "fi": "Finnish", "ru": "Russian",
        "zh": "Chinese", "ja": "Japanese", "ko": "Korean", "ar": "Arabic"
    }
    language_display = language_names.get(language_code, language_code.upper())
    
    metadata = {
        'interviewer': interviewer_name,
        'participant': participant_name,
        'date': datetime.datetime.now().strftime('%Y-%m-%d'),
        'duration': format_duration(audio_duration),
        'language': language_display,
        'interview_type': "Research interview",  # Generic description
        'audio_quality': get_audio_quality_estimate(result),
        'total_segments': len(segments),
        'total_words': total_words,
        'avg_words_per_segment': round(avg_words_per_segment, 1),
        'avg_confidence': round(avg_confidence, 3),
        'high_confidence_pct': round((high_confidence_segments / len(segments)) * 100, 1),
        'medium_confidence_pct': round((medium_confidence_segments / len(segments)) * 100, 1),
        'low_confidence_pct': round((low_confidence_segments / len(segments)) * 100, 1),
        'avg_segment_length': round(avg_segment_length, 2),
        'longest_segment': round(longest_segment, 2),
        'shortest_segment': round(shortest_segment, 2),
        'total_pauses': len(pauses),
        'avg_pause_duration': round(avg_pause_duration, 2),
        'longest_pause': round(longest_pause, 2),
        'total_pause_time': round(total_pause_time, 2),
        'speech_vs_pause_ratio': round((audio_duration - total_pause_time) / audio_duration * 100, 1) if audio_duration > 0 else 0
    }
    
    # Process segments with enhanced speaker detection
    segments = result["segments"]
    speakers = detect_speaker_changes(segments, silence_threshold=1.2, min_speaker_time=4.0)
    
    # Create transcript data
    transcript_segments = []
    for i, segment in enumerate(segments):
        segment_data = {
            'start_time': format_timestamp(segment['start']),
            'end_time': format_timestamp(segment['end']),
            'speaker': speakers[i],
            'text': segment['text'].strip(),
            'confidence': segment.get('avg_logprob', 0)
        }
        transcript_segments.append(segment_data)
    
    # Print to console
    print("\n" + "="*80)
    print("ACADEMIC INTERVIEW TRANSCRIPT")
    print("="*80)
    print(f"Study Date: {metadata['date']}")
    print(f"Duration: {metadata['duration']}")
    print(f"Language: {metadata['language']}")
    print(f"Quality: {metadata['audio_quality']}")
    print("="*80 + "\n")
    
    current_speaker = None
    turn_number = 0
    
    for segment in transcript_segments:
        if current_speaker != segment['speaker']:
            if current_speaker is not None:
                print()
            current_speaker = segment['speaker']
            turn_number += 1
        
        marked_text = add_content_markers(segment['text'])
        print(f"Turn {turn_number:2d} [{segment['start_time']} - {segment['end_time']}] {segment['speaker']}: {marked_text}")
    
    # Save files to project folder
    print(f"\nSaving files to project folder...")
    
    # Text file
    text_file_path = os.path.join(project_folder, "academic_transcript.txt")
    with open(text_file_path, 'w', encoding='utf-8') as f:
        f.write("ACADEMIC INTERVIEW TRANSCRIPT\n")
        f.write("=" * 50 + "\n\n")
        f.write(f"Study Date: {metadata['date']}\n")
        f.write(f"Duration: {metadata['duration']}\n") 
        f.write(f"Language: {metadata['language']}\n")
        f.write(f"Quality: {metadata['audio_quality']}\n")
        f.write("=" * 50 + "\n\n")
        
        current_speaker = None
        turn_number = 0
        
        for segment in transcript_segments:
            if current_speaker != segment['speaker']:
                if current_speaker is not None:
                    f.write("\n")
                current_speaker = segment['speaker']
                turn_number += 1
            
            marked_text = add_content_markers(segment['text'])
            f.write(f"Turn {turn_number:2d} [{segment['start_time']} - {segment['end_time']}] {segment['speaker']}: {marked_text}\n")
    
    print(f"Text transcript saved: {text_file_path}")
    
    # Word document
    try:
        word_file_path = os.path.join(project_folder, "academic_transcript.docx")
        create_academic_word_document(transcript_segments, metadata, word_file_path)
    except ImportError:
        print("Note: python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        print(f"Error creating Word document: {e}")
    
    # Summary with enhanced statistics
    print(f"\n" + "="*80)
    print("TRANSCRIPTION SUMMARY")
    print("="*80)
    print(f"Project folder: {project_folder}")
    print(f"Interviewer: {metadata['interviewer']}")
    print(f"Participant: {metadata['participant']}")
    print(f"Language: {metadata['language']}")
    print(f"Total duration: {metadata['duration']}")
    print(f"Total turns: {len(transcript_segments)}")
    print(f"Total words: {metadata['total_words']}")
    
    # Calculate speech rate
    duration_parts = metadata['duration'].split(':')
    total_minutes = int(duration_parts[0]) + int(duration_parts[1])/60
    speech_rate = metadata['total_words'] / total_minutes if total_minutes > 0 else 0
    print(f"Speech rate: {speech_rate:.1f} words/minute")
    
    speaker_counts = {}
    speaker_word_counts = {}
    for segment in transcript_segments:
        speaker = segment['speaker']
        speaker_counts[speaker] = speaker_counts.get(speaker, 0) + 1
        word_count = len(segment['text'].split())
        speaker_word_counts[speaker] = speaker_word_counts.get(speaker, 0) + word_count
    
    for speaker, count in speaker_counts.items():
        percentage = (count / len(transcript_segments)) * 100
        word_percentage = (speaker_word_counts[speaker] / metadata['total_words']) * 100
        print(f"{speaker}: {count} turns ({percentage:.1f}%), {speaker_word_counts[speaker]} words ({word_percentage:.1f}%)")
    
    print(f"\nQuality Metrics:")
    print(f"Audio quality: {metadata['audio_quality']}")
    print(f"Average confidence: {metadata['avg_confidence']} (higher is better)")
    print(f"High confidence segments: {metadata['high_confidence_pct']}%")
    print(f"Speech vs silence ratio: {metadata['speech_vs_pause_ratio']}%")
    print(f"Average segment length: {metadata['avg_segment_length']}s")
    print(f"Average pause duration: {metadata['avg_pause_duration']}s")
    
    print(f"\nFiles created: audio copy, .txt, .docx")
    print(f"Context provided to model: Yes ({metadata['interview_type']} context)")
    print(f"Enhanced quality parameters: Enabled (beam search, best-of-5)")

if __name__ == "__main__":
    main()